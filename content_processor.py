"""
Content processor for converting uploaded files to clean Markdown.

Supports: .md, .pdf, .txt, .docx
All content is normalized to Markdown for consistent storage and AI processing.
"""

import os
import io
from typing import Tuple


def process_file(uploaded_file) -> Tuple[str, bytes]:
    """
    Process an uploaded file and return (cleaned_markdown, original_bytes).
    Determines file type by extension and delegates to the appropriate converter.
    """
    filename = uploaded_file.name
    ext = os.path.splitext(filename)[1].lower()
    raw_bytes = uploaded_file.read()

    if ext == ".md":
        markdown = raw_bytes.decode("utf-8", errors="replace")
    elif ext == ".txt":
        markdown = _txt_to_markdown(raw_bytes)
    elif ext == ".pdf":
        markdown = _pdf_to_markdown(raw_bytes)
    elif ext == ".docx":
        markdown = _docx_to_markdown(raw_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Clean up excessive whitespace
    lines = markdown.splitlines()
    cleaned = "\n".join(line.rstrip() for line in lines)
    # Collapse 3+ blank lines into 2
    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")
    cleaned = cleaned.strip()

    return cleaned, raw_bytes


def _txt_to_markdown(raw_bytes: bytes) -> str:
    """Convert plain text to markdown by preserving paragraphs."""
    text = raw_bytes.decode("utf-8", errors="replace")
    return text


def _pdf_to_markdown(raw_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF and format as markdown."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return "[Error: PyMuPDF not installed. Run: pip install PyMuPDF]"

    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"## Page {i + 1}\n\n{text.strip()}")
    doc.close()
    return "\n\n".join(pages)


def _docx_to_markdown(raw_bytes: bytes) -> str:
    """Convert DOCX to markdown using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return "[Error: python-docx not installed. Run: pip install python-docx]"

    doc = Document(io.BytesIO(raw_bytes))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""
        if "heading 1" in style:
            parts.append(f"# {text}")
        elif "heading 2" in style:
            parts.append(f"## {text}")
        elif "heading 3" in style:
            parts.append(f"### {text}")
        elif "list" in style:
            parts.append(f"- {text}")
        else:
            parts.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            if i == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        parts.append("\n".join(rows))

    return "\n\n".join(parts)
